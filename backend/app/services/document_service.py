import os
import uuid
import fitz  # PyMuPDF
import docx
import httpx
from typing import Optional, List, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from fastapi import HTTPException, UploadFile
from app.models.document import Document
from app.core.config import settings
from app.services.text_extractor import TextExtractor
from app.services.claim_extractor import ClaimExtractor
from app.services.quality_scorer import QualityScorer


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.text_extractor = TextExtractor()
        self.claim_extractor = ClaimExtractor()
        self.quality_scorer = QualityScorer()

    async def process_uploaded_file(self, file: UploadFile) -> Document:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")

        file_extension = file.filename.split(".")[-1].lower()
        if file_extension not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed. Allowed types: {settings.ALLOWED_EXTENSIONS}",
            )

        # Save uploaded file
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        file_id = str(uuid.uuid4())
        file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}.{file_extension}")

        content = await file.read()
        if len(content) > settings.MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="File too large")

        with open(file_path, "wb") as f:
            f.write(content)

        # Extract text and metadata
        extracted_data = await self.text_extractor.extract_from_file(
            file_path, file_extension
        )

        # Extract claims
        claims = await self.claim_extractor.extract_claims(extracted_data["content"])

        # Calculate quality score
        quality_score = await self.quality_scorer.calculate_score(
            extracted_data["content"], claims
        )

        # Create document record
        document = Document(
            title=extracted_data.get("title", file.filename),
            authors=extracted_data.get("authors", []),
            abstract=extracted_data.get("abstract", ""),
            content=extracted_data["content"],
            file_path=file_path,
            file_type=file_extension,
            extracted_claims=claims,
            method_quality_score=quality_score,
            confidence_score=quality_score * 0.8,  # Initial confidence based on quality
        )

        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)

        return document

    async def process_url(self, url: str) -> Document:
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, timeout=30)
                response.raise_for_status()

            # Try to extract PDF content if it's a PDF URL
            if url.lower().endswith(
                ".pdf"
            ) or "application/pdf" in response.headers.get("content-type", ""):
                # Save as temporary PDF file
                file_id = str(uuid.uuid4())
                file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}.pdf")
                os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

                with open(file_path, "wb") as f:
                    f.write(response.content)

                extracted_data = await self.text_extractor.extract_from_file(
                    file_path, "pdf"
                )
            else:
                # Extract from HTML
                extracted_data = await self.text_extractor.extract_from_html(
                    response.text, url
                )
                file_path = None

            # Extract claims
            claims = await self.claim_extractor.extract_claims(
                extracted_data["content"]
            )

            # Calculate quality score
            quality_score = await self.quality_scorer.calculate_score(
                extracted_data["content"], claims
            )

            # Create document record
            document = Document(
                title=extracted_data.get("title", "Document from URL"),
                authors=extracted_data.get("authors", []),
                abstract=extracted_data.get("abstract", ""),
                content=extracted_data["content"],
                url=url,
                file_path=file_path,
                file_type="url",
                extracted_claims=claims,
                method_quality_score=quality_score,
                confidence_score=quality_score * 0.8,
            )

            self.db.add(document)
            await self.db.commit()
            await self.db.refresh(document)

            return document

        except httpx.RequestError as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to fetch URL: {str(e)}"
            )

    async def process_doi(self, doi: str) -> Document:
        # Check if document already exists
        result = await self.db.execute(select(Document).where(Document.doi == doi))
        existing_doc = result.scalar_one_or_none()
        if existing_doc:
            return existing_doc

        # Fetch from DOI
        doi_url = f"https://doi.org/{doi}"
        try:
            # First try to get metadata from Crossref
            crossref_url = f"https://api.crossref.org/works/{doi}"
            async with httpx.AsyncClient() as client:
                response = await client.get(crossref_url, timeout=30)

            if response.status_code == 200:
                data = response.json()
                work = data["message"]

                title = work.get("title", ["Untitled"])[0]
                authors = [
                    f"{author.get('given', '')} {author.get('family', '')}"
                    for author in work.get("author", [])
                ]
                abstract = work.get("abstract", "")

                # Try to get full text from DOI redirect
                full_text = ""
                try:
                    async with httpx.AsyncClient() as client:
                        pdf_response = await client.get(
                            doi_url, timeout=30, follow_redirects=True
                        )
                    if "application/pdf" in pdf_response.headers.get(
                        "content-type", ""
                    ):
                        # Save and extract PDF
                        file_id = str(uuid.uuid4())
                        file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}.pdf")
                        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

                        with open(file_path, "wb") as f:
                            f.write(pdf_response.content)

                        extracted_data = await self.text_extractor.extract_from_file(
                            file_path, "pdf"
                        )
                        full_text = extracted_data["content"]
                    else:
                        # Try HTML extraction
                        extracted_data = await self.text_extractor.extract_from_html(
                            pdf_response.text, doi_url
                        )
                        full_text = extracted_data["content"]
                        file_path = None
                except:
                    # Fallback to metadata only
                    full_text = f"{title}\n\n{abstract}"
                    file_path = None

                # Extract claims
                claims = await self.claim_extractor.extract_claims(full_text)

                # Calculate quality score
                quality_score = await self.quality_scorer.calculate_score(
                    full_text, claims
                )

                # Create document record
                document = Document(
                    title=title,
                    authors=authors,
                    abstract=abstract,
                    content=full_text,
                    doi=doi,
                    url=doi_url,
                    file_path=file_path,
                    file_type="doi",
                    extracted_claims=claims,
                    method_quality_score=quality_score,
                    confidence_score=quality_score * 0.8,
                )

                self.db.add(document)
                await self.db.commit()
                await self.db.refresh(document)

                return document

        except httpx.RequestError as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to fetch DOI: {str(e)}"
            )

        raise HTTPException(status_code=404, detail="DOI not found")

    async def get_document(self, document_id: int) -> Optional[Document]:
        result = await self.db.execute(
            select(Document).where(Document.id == document_id)
        )
        return result.scalar_one_or_none()

    async def list_documents(self, limit: int = 50) -> List[Document]:
        result = await self.db.execute(
            select(Document).order_by(Document.created_at.desc()).limit(limit)
        )
        return list(result.scalars().all())
