import re
import fitz  # PyMuPDF
import docx
from typing import Dict, List, Any
from bs4 import BeautifulSoup

class TextExtractor:
    def __init__(self):
        pass
    
    async def extract_from_file(self, file_path: str, file_type: str) -> Dict[str, Any]:
        if file_type == "pdf":
            return await self._extract_from_pdf(file_path)
        elif file_type == "docx":
            return await self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        doc = fitz.open(file_path)
        
        # Extract metadata
        metadata = doc.metadata
        title = metadata.get("title", "")
        authors = self._parse_authors(metadata.get("author", ""))
        
        # Extract text content
        full_text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            full_text += page.get_text()
        
        doc.close()
        
        # Extract abstract
        abstract = self._extract_abstract(full_text)
        
        # Clean and structure text
        structured_content = self._structure_content(full_text)
        
        return {
            "title": title or self._extract_title(full_text),
            "authors": authors,
            "abstract": abstract,
            "content": structured_content,
            "raw_text": full_text
        }
    
    async def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        doc = docx.Document(file_path)
        
        # Extract core properties
        props = doc.core_properties
        title = props.title or ""
        authors = self._parse_authors(props.author or "")
        
        # Extract text content
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + " "
                full_text += "\n"
        
        # Extract abstract
        abstract = self._extract_abstract(full_text)
        
        # Clean and structure text
        structured_content = self._structure_content(full_text)
        
        return {
            "title": title or self._extract_title(full_text),
            "authors": authors,
            "abstract": abstract,
            "content": structured_content,
            "raw_text": full_text
        }
    
    async def extract_from_html(self, html_content: str, url: str) -> Dict[str, Any]:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract title
        title_tag = soup.find('title')
        title = title_tag.get_text().strip() if title_tag else ""
        
        # Try to find title in meta tags or h1
        if not title:
            meta_title = soup.find('meta', property='og:title') or soup.find('meta', name='title')
            if meta_title:
                title = meta_title.get('content', '')
        
        if not title:
            h1 = soup.find('h1')
            if h1:
                title = h1.get_text().strip()
        
        # Extract authors from meta tags
        authors = []
        author_tags = soup.find_all('meta', name='author') or soup.find_all('meta', property='article:author')
        for tag in author_tags:
            author = tag.get('content', '')
            if author:
                authors.extend(self._parse_authors(author))
        
        # Extract main content
        # Remove script, style, and other non-content elements
        for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
            element.decompose()
        
        # Try to find main content area
        main_content = (
            soup.find('main') or 
            soup.find('article') or 
            soup.find('div', class_=re.compile(r'content|article|main', re.I)) or
            soup.find('body')
        )
        
        if main_content:
            full_text = main_content.get_text(separator='\n', strip=True)
        else:
            full_text = soup.get_text(separator='\n', strip=True)
        
        # Extract abstract
        abstract = self._extract_abstract(full_text)
        
        # Clean and structure text
        structured_content = self._structure_content(full_text)
        
        return {
            "title": title,
            "authors": authors,
            "abstract": abstract,
            "content": structured_content,
            "raw_text": full_text,
            "url": url
        }
    
    def _parse_authors(self, author_string: str) -> List[str]:
        if not author_string:
            return []
        
        # Common author separators
        separators = [';', ',', ' and ', ' & ', '\n']
        authors = [author_string]
        
        for sep in separators:
            new_authors = []
            for author in authors:
                new_authors.extend([a.strip() for a in author.split(sep) if a.strip()])
            authors = new_authors
        
        # Clean up author names
        cleaned_authors = []
        for author in authors:
            # Remove common prefixes/suffixes
            author = re.sub(r'^(Dr\.?|Prof\.?|Mr\.?|Ms\.?|Mrs\.?)\s+', '', author, flags=re.I)
            author = re.sub(r'\s+(Ph\.?D\.?|M\.?D\.?|Jr\.?|Sr\.?)$', '', author, flags=re.I)
            
            if len(author) > 2:  # Filter out initials and short strings
                cleaned_authors.append(author)
        
        return cleaned_authors[:10]  # Limit to reasonable number
    
    def _extract_title(self, text: str) -> str:
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 200:  # Reasonable title length
                # Check if it looks like a title (not all caps, has words)
                if not line.isupper() and len(line.split()) > 2:
                    return line
        return "Untitled Document"
    
    def _extract_abstract(self, text: str) -> str:
        # Look for abstract section
        abstract_patterns = [
            r'(?i)abstract\s*[:\-]?\s*(.+?)(?=\n\s*\n|\n\s*(?:introduction|keywords|1\.|background))',
            r'(?i)summary\s*[:\-]?\s*(.+?)(?=\n\s*\n|\n\s*(?:introduction|keywords|1\.|background))',
        ]
        
        for pattern in abstract_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                abstract = match.group(1).strip()
                # Clean up the abstract
                abstract = re.sub(r'\s+', ' ', abstract)
                if 50 < len(abstract) < 2000:  # Reasonable abstract length
                    return abstract
        
        # Fallback: take first paragraph that's long enough
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if 100 < len(para) < 1000:
                return para
        
        return ""
    
    def _structure_content(self, text: str) -> str:
        # Basic text cleaning and structuring
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'\n\s*Page \d+.*?\n', '\n', text, flags=re.I)
        
        return text.strip()